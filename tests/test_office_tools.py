"""
Edge case tests for tools/office_tools.py.
All fixtures are built programmatically — no binary files committed to the repo.
"""
import os
import sys
import pytest
sys.path.insert(0, os.path.join(os.path.dirname(__file__), '..', 'tools-harness'))


# ── Fixtures ────────────────────────────────────────────────────────────────────

def _make_xlsx(path, sheets: dict):
    """Create an xlsx file at path. sheets = {sheet_name: [[row], ...]}"""
    import openpyxl
    wb = openpyxl.Workbook()
    first = True
    for name, rows in sheets.items():
        ws = wb.active if first else wb.create_sheet()
        ws.title = name
        for row in rows:
            ws.append(row)
        first = False
    wb.save(str(path))


def _make_docx(path, paragraphs: list, heading_style: str | None = None):
    """Create a .docx file at path with the given paragraphs."""
    import docx
    doc = docx.Document()
    for i, text in enumerate(paragraphs):
        if i == 0 and heading_style:
            doc.add_heading(text, level=int(heading_style[-1]))
        else:
            doc.add_paragraph(text)
    doc.save(str(path))


def _make_docx_with_table(path, rows: list[list[str]]):
    import docx
    doc = docx.Document()
    if rows:
        table = doc.add_table(rows=len(rows), cols=len(rows[0]))
        for r_idx, row in enumerate(rows):
            for c_idx, val in enumerate(row):
                table.rows[r_idx].cells[c_idx].text = val
    doc.save(str(path))


# ── Excel: happy path ───────────────────────────────────────────────────────────

def test_excel_to_markdown_basic(tmp_path):
    from tools.office_tools import excel_to_markdown
    xls = tmp_path / "data.xlsx"
    _make_xlsx(xls, {"Sheet1": [["Name", "Score"], ["Alice", 95], ["Bob", 87]]})

    md_path, content = excel_to_markdown(str(xls), str(tmp_path))
    assert md_path.endswith(".md")
    assert os.path.exists(md_path)
    assert "Sheet1" in content
    assert "Alice" in content
    assert "95" in content


def test_excel_to_markdown_multiple_sheets(tmp_path):
    from tools.office_tools import excel_to_markdown
    xls = tmp_path / "multi.xlsx"
    _make_xlsx(xls, {
        "Q1": [["Revenue", "1000"]],
        "Q2": [["Revenue", "2000"]],
    })
    _, content = excel_to_markdown(str(xls), str(tmp_path))
    assert "Q1" in content
    assert "Q2" in content


def test_excel_to_markdown_empty_sheet(tmp_path):
    from tools.office_tools import excel_to_markdown
    xls = tmp_path / "empty.xlsx"
    _make_xlsx(xls, {"EmptySheet": []})
    _, content = excel_to_markdown(str(xls), str(tmp_path))
    assert "EmptySheet" in content
    assert "empty sheet" in content.lower()


def test_excel_to_markdown_none_cells(tmp_path):
    """Rows with None cells must not crash — they should render as empty string."""
    from tools.office_tools import excel_to_markdown
    import openpyxl
    xls = tmp_path / "nones.xlsx"
    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = "Data"
    ws.append(["A", None, "C"])
    ws.append([None, None, None])
    ws.append([1, None, 3])
    wb.save(str(xls))

    _, content = excel_to_markdown(str(xls), str(tmp_path))
    assert "A" in content
    assert "C" in content


def test_excel_to_markdown_large_sheet_truncated(tmp_path):
    """Sheets with >200 rows must be truncated without crashing."""
    from tools.office_tools import excel_to_markdown
    xls = tmp_path / "big.xlsx"
    rows = [[f"row{i}", i] for i in range(500)]
    _make_xlsx(xls, {"BigSheet": rows})

    _, content = excel_to_markdown(str(xls), str(tmp_path))
    # row0..row199 present, row200+ absent (cap is 200 rows)
    assert "row0" in content
    assert "row199" in content
    assert "row200" not in content


def test_excel_to_markdown_missing_file(tmp_path):
    from tools.office_tools import excel_to_markdown
    with pytest.raises(FileNotFoundError):
        excel_to_markdown(str(tmp_path / "nonexistent.xlsx"), str(tmp_path))


# ── Excel metadata ──────────────────────────────────────────────────────────────

def test_excel_metadata_returns_sheet_row_counts(tmp_path):
    from tools.office_tools import excel_metadata
    xls = tmp_path / "meta.xlsx"
    _make_xlsx(xls, {
        "Alpha": [["a"], ["b"], ["c"]],
        "Beta": [["x"]],
    })
    meta = excel_metadata(str(xls))
    assert set(meta.keys()) == {"Alpha", "Beta"}
    assert meta["Alpha"] == 3
    assert meta["Beta"] == 1


# ── Word: happy path ────────────────────────────────────────────────────────────

def test_docx_to_markdown_basic(tmp_path):
    from tools.office_tools import docx_to_markdown
    doc = tmp_path / "letter.docx"
    _make_docx(doc, ["Hello world", "Second paragraph"])

    md_path, content = docx_to_markdown(str(doc), str(tmp_path))
    assert md_path.endswith(".md")
    assert os.path.exists(md_path)
    assert "Hello world" in content
    assert "Second paragraph" in content


def test_docx_to_markdown_heading_styles(tmp_path):
    """Heading 1 paragraphs must render as # in output."""
    import docx as _docx
    from tools.office_tools import docx_to_markdown
    doc_path = tmp_path / "headed.docx"
    doc = _docx.Document()
    doc.add_heading("Main Title", level=1)
    doc.add_heading("Sub Section", level=2)
    doc.add_paragraph("Body text here.")
    doc.save(str(doc_path))

    _, content = docx_to_markdown(str(doc_path), str(tmp_path))
    assert "# Main Title" in content
    assert "## Sub Section" in content
    assert "Body text here." in content


def test_docx_to_markdown_table(tmp_path):
    from tools.office_tools import docx_to_markdown
    doc_path = tmp_path / "table.docx"
    _make_docx_with_table(doc_path, [
        ["Name", "Dept", "Score"],
        ["Alice", "Eng", "95"],
        ["Bob", "HR", "88"],
    ])
    _, content = docx_to_markdown(str(doc_path), str(tmp_path))
    assert "Name" in content
    assert "Alice" in content
    assert "|" in content


def test_docx_to_markdown_empty_document(tmp_path):
    """Empty .docx must not crash — returns a file with just the heading."""
    import docx as _docx
    from tools.office_tools import docx_to_markdown
    doc_path = tmp_path / "empty.docx"
    _docx.Document().save(str(doc_path))

    md_path, content = docx_to_markdown(str(doc_path), str(tmp_path))
    assert os.path.exists(md_path)
    assert isinstance(content, str)


def test_docx_to_markdown_missing_file(tmp_path):
    from tools.office_tools import docx_to_markdown
    with pytest.raises(FileNotFoundError):
        docx_to_markdown(str(tmp_path / "ghost.docx"), str(tmp_path))


def test_docx_to_markdown_only_whitespace_paragraphs(tmp_path):
    """Paragraphs with only whitespace must be skipped, not appended as blank lines."""
    import docx as _docx
    from tools.office_tools import docx_to_markdown
    doc_path = tmp_path / "ws.docx"
    doc = _docx.Document()
    doc.add_paragraph("   ")
    doc.add_paragraph("")
    doc.add_paragraph("Real content")
    doc.save(str(doc_path))

    _, content = docx_to_markdown(str(doc_path), str(tmp_path))
    assert "Real content" in content
    # The blank/whitespace paragraphs should not produce runs of empty lines
    assert "\n\n\n" not in content
