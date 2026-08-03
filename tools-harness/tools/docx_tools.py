"""DOCX form processing utilities: detect form fields and fill them."""

from __future__ import annotations

import json
from pathlib import Path

from docx import Document


def detect_docx_form_fields(path: str) -> str:
    """Inspect a DOCX file for form fields (text boxes in tables/structured content).

    Args:
        path: Absolute or relative path to the DOCX file.

    Returns:
        A JSON-formatted string listing every form field found, or a plain message
        if the DOCX has no obvious form structure.

    Raises:
        FileNotFoundError: If the DOCX does not exist.
    """
    p = Path(path)
    if not p.exists():
        raise FileNotFoundError(f"DOCX not found: {path}")

    doc = Document(str(p))
    fields: list[dict] = []
    field_id = 1

    # Strategy: scan tables for labeled cells (label in one cell, blank cell for input)
    # Also detect any paragraphs with specific patterns (colon-terminated labels)
    for table_idx, table in enumerate(doc.tables):
        for row_idx, row in enumerate(table.rows):
            for cell_idx, cell in enumerate(row.cells):
                text = cell.text.strip()
                # Detect label patterns: "Full Name:", "Email:", etc.
                if text.endswith(":") and len(text) < 100:
                    # This looks like a form field label — value lives in the next cell
                    # of the same row (label/value side-by-side table layout).
                    value = ""
                    if cell_idx + 1 < len(row.cells):
                        value = row.cells[cell_idx + 1].text.strip()
                    field_name = text.rstrip(":").strip().replace(" ", "_").lower()
                    fields.append({
                        "id": field_id,
                        "name": field_name,
                        "label": text.rstrip(":"),
                        "type": "text",
                        "location": f"table_{table_idx}_row_{row_idx}_cell_{cell_idx}",
                        "current_value": value,
                    })
                    field_id += 1

    # Also scan body paragraphs for label patterns
    for para_idx, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        if text.endswith(":") and 10 < len(text) < 100:
            field_name = text.rstrip(":").strip().replace(" ", "_").lower()
            if not any(f["label"] == text.rstrip(":") for f in fields):
                fields.append({
                    "id": field_id,
                    "name": field_name,
                    "label": text.rstrip(":"),
                    "type": "text",
                    "location": f"para_{para_idx}",
                    "current_value": "",
                })
                field_id += 1

    if not fields:
        return (
            "This DOCX contains no obvious form fields. "
            "Tip: Form fields are typically labeled cells (ending with ':') in tables."
        )

    return json.dumps({"fields": fields, "total": len(fields)}, indent=2, ensure_ascii=False)


def fill_docx_form(
    path: str,
    fields: dict | str,
    output_path: str | None = None,
) -> str:
    """Fill form fields in a DOCX and save to a new file.

    Args:
        path: Absolute or relative path to the source DOCX.
        fields: A mapping of {field_name: value} or {label: value}. May be passed as
                a JSON string (executor handles parse, but function accepts both).
        output_path: Where to write the filled DOCX. Defaults to
                     ``{stem}_filled.docx`` in the same directory as the source.

    Returns:
        A plain-text confirmation message listing filled fields.

    Raises:
        FileNotFoundError: If the source DOCX does not exist.
        ValueError: If ``fields`` is a string that cannot be parsed as JSON.
    """
    p = Path(path)
    if not p.exists():
        raise FileNotFoundError(f"DOCX not found: {path}")

    # Accept JSON string or dict
    if isinstance(fields, str):
        try:
            fields = json.loads(fields)
        except json.JSONDecodeError as exc:
            raise ValueError(f"fields must be a JSON object string or dict: {exc}") from exc

    doc = Document(str(p))
    out = Path(output_path) if output_path else p.parent / f"{p.stem}_filled.docx"

    filled: list[str] = []
    skipped: list[str] = []

    # Normalize field keys to lowercase underscored versions for matching
    normalized_fields = {k.replace(" ", "_").lower(): v for k, v in fields.items()}

    # Fill table cells (strategy: label in cell N, value in cell N+1 or next row)
    for table in doc.tables:
        for row_idx, row in enumerate(table.rows):
            for cell_idx, cell in enumerate(row.cells):
                text = cell.text.strip()
                # Check if this is a label cell
                if text.endswith(":"):
                    field_name = text.rstrip(":").strip().replace(" ", "_").lower()
                    if field_name in normalized_fields:
                        # Try to fill the next cell in the same row
                        if cell_idx + 1 < len(row.cells):
                            target_cell = row.cells[cell_idx + 1]
                            target_cell.paragraphs[0].clear()
                            target_cell.paragraphs[0].text = str(normalized_fields[field_name])
                            filled.append(field_name)
                        # Or the same cell's following paragraph
                        elif len(cell.paragraphs) > 1:
                            cell.paragraphs[1].clear()
                            cell.paragraphs[1].text = str(normalized_fields[field_name])
                            filled.append(field_name)

    # Fill body paragraphs (in-paragraph filling)
    for para in doc.paragraphs:
        text = para.text.strip()
        if text.endswith(":"):
            field_name = text.rstrip(":").strip().replace(" ", "_").lower()
            if field_name in normalized_fields:
                para.clear()
                para.text = f"{text} {normalized_fields[field_name]}"
                filled.append(field_name)

    # Track skipped (requested but not found)
    for field_name in normalized_fields:
        if field_name not in filled:
            skipped.append(field_name)

    doc.save(str(out))

    lines = [f"Saved filled DOCX → {out}"]
    if filled:
        lines.append(f"Fields filled ({len(filled)}): {', '.join(filled)}")
    if skipped:
        lines.append(f"Fields skipped — not found in DOCX ({len(skipped)}): {', '.join(skipped)}")
    return "\n".join(lines)
